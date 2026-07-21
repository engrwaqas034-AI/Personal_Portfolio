import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900 / Navy
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="0D9488"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def create_cv_docx(output_path, img_path):
    doc = docx.Document()

    # Page Margins (0.5 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)

    # ==========================================
    # HEADER SECTION (Table with Text & Profile Photo)
    # ==========================================
    tbl_header = doc.add_table(rows=1, cols=2)
    tbl_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_header.autofit = False

    cell_text = tbl_header.rows[0].cells[0]
    cell_img = tbl_header.rows[0].cells[1]

    cell_text.width = Inches(5.8)
    cell_img.width = Inches(1.3)

    set_cell_margins(cell_text, top=0, bottom=0, left=0, right=50)
    set_cell_margins(cell_img, top=0, bottom=0, left=50, right=0)
    
    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Fill Text Cell
    p_name = cell_text.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run("MUHAMMAD WAQAS")
    r_name.font.size = Pt(21)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(15, 23, 42)

    p_title = cell_text.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("CIVIL ENGINEERING  |  BIM & VDC SPECIALIST  |  APPLIED AI & LLM ENGINEER")
    r_title.font.size = Pt(9.5)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(13, 148, 136)

    p_contact = cell_text.add_paragraph()
    p_contact.paragraph_format.space_after = Pt(2)
    r_contact = p_contact.add_run("Rawalpindi, Punjab, Pakistan   |   +92 345 5862998   |   engrwaqas034@gmail.com")
    r_contact.font.size = Pt(9)
    r_contact.font.color.rgb = RGBColor(71, 85, 105)

    p_links = cell_text.add_paragraph()
    p_links.paragraph_format.space_after = Pt(6)
    
    r_l1 = p_links.add_run("LinkedIn: ")
    r_l1.font.bold = True
    r_l1.font.size = Pt(8.5)
    p_links.add_run("linkedin.com/in/muhammad-waqas-27b232214   |   ")
    
    r_l2 = p_links.add_run("GitHub: ")
    r_l2.font.bold = True
    r_l2.font.size = Pt(8.5)
    p_links.add_run("github.com/engrwaqas034-AI")

    # Add Photo to Image Cell
    if os.path.exists(img_path):
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(img_path, width=Inches(1.25))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ==========================================
    # EXECUTIVE SUMMARY
    # ==========================================
    add_heading_styled(doc, "EXECUTIVE PROFILE")
    
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(10)
    p_summary.paragraph_format.line_spacing = 1.15
    r_sum = p_summary.add_run(
        "Results-driven Civil Engineering Specialist, Manager Estate (Technical) at REDAMCO (Ministry of Railways), and Senior Infrastructure Professional "
        "with 15+ years of field execution, project management, and technical appraisal experience across public sector mega-railways (Pakistan Railways), "
        "GCC/Oman projects, and multi-story housing developments. Certified ISO 19650 Information Manager & BIM Specialist proficient in Revit, "
        "Navisworks Manage (4D/5D), Civil 3D, Primavera P6, and Autodesk Construction Cloud. Applied AI & LLM Engineer specializing in AI Agentic workflows (N8N), "
        "generative architectural engines (AI-Architect), construction claim RAG intelligence platforms, and quantity surveying automation. Proven expertise "
        "in land monetization, project feasibility analysis, business development support, and digitally-enabled infrastructure delivery."
    )
    r_sum.font.size = Pt(9.5)

    # ==========================================
    # CORE COMPETENCIES & TECHNICAL MATRIX
    # ==========================================
    add_heading_styled(doc, "CORE COMPETENCIES & TECHNICAL MATRIX")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    col_widths = [Inches(3.5), Inches(3.5)]
    
    matrix_data = [
        ("Estate Management & Infrastructure", [
            "Technical Project Appraisal & Execution Analysis",
            "Business Development & Asset Monetization Support",
            "Railway Land Development & Commercial Feasibilities",
            "Railway Infrastructure & Track Maintenance",
            "GIS Land Surveying & Encroachment Removal",
            "Quantity Surveying, IPCs & Contractual Claims"
        ]),
        ("BIM, VDC & ISO 19650 Management", [
            "Autodesk Revit (Architectural & Structural)",
            "Navisworks Manage (Clash & 4D/5D Simulation)",
            "Civil 3D (Corridor Design & Cut/Fill Earthworks)",
            "Primavera P6 (Scheduling, Baseline, EVM)",
            "ISO 19650 Information Management (Plannerly Certified)",
            "ACC / BIM 360, ReCap Pro, Twinmotion, ETABS"
        ]),
        ("Applied AI & AEC Automation", [
            "AI Architectural Layout Engine (AI-Architect)",
            "Construction Claim Analyzer (LLM + RAG)",
            "AI BOQ & Quantity Surveying Generator",
            "N8N AI Agentic Workflow Automation",
            "Python for AEC Systems & API Integration",
            "Machine Learning & Algorithmic Design"
        ]),
        ("Software & Web Development", [
            "Python, HTML5, CSS3, JavaScript",
            "Git / GitHub Version Control & Repositories",
            "Rest API Development & Web Dashboards",
            "Full-Stack Web App Development",
            "Microsoft Project & Advanced MS Office",
            "Technical Reporting & Project Governance"
        ])
    ]

    idx = 0
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_widths[idx % 2]
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            title, items = matrix_data[idx]
            p_cat = cell.paragraphs[0]
            p_cat.paragraph_format.space_after = Pt(4)
            r_cat = p_cat.add_run(title)
            r_cat.font.bold = True
            r_cat.font.size = Pt(9.5)
            r_cat.font.color.rgb = RGBColor(13, 148, 136)

            for item in items:
                p_item = cell.add_paragraph()
                p_item.paragraph_format.space_after = Pt(2)
                p_item.paragraph_format.left_indent = Inches(0.15)
                r_bullet = p_item.add_run("▪  ")
                r_bullet.font.size = Pt(8)
                r_bullet.font.color.rgb = RGBColor(13, 148, 136)
                r_txt = p_item.add_run(item)
                r_txt.font.size = Pt(8.5)

            idx += 1

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ==========================================
    # FEATURED AI & DIGITAL AEC PROJECTS
    # ==========================================
    add_heading_styled(doc, "FEATURED AI & DIGITAL AEC INNOVATION PROJECTS")

    ai_projects = [
        ("AI-Architect — Generative AI Floorplan & CAD/BIM Engine", 
         "GitHub Repo: github.com/engrwaqas034-AI/AI-Architect",
         "Architectural & AI Software Project",
         [
             "Developed an end-to-end Python & Web AI engine that automates 2D spatial layouts, boundary envelope math, setback compliance, and floor plan generation.",
             "Integrated real-time parametric SVG rendering, multi-floor space allocations, and direct export pipelines for CAD/BIM workflows.",
             "Designed custom CSP (Constraint Satisfaction Problem) solvers to eliminate room overlap and optimize natural lighting and ventilation."
         ]),
        ("Construction_Claim_Analyzer — LLM + RAG Claim Intelligence System",
         "GitHub Repo: github.com/engrwaqas034-AI/Construction_Claim_Analyzer",
         "AI & Legal/Contractual Engineering Project",
         [
             "Engineered an LLM and Retrieval-Augmented Generation (RAG) system specialized in evaluating construction delay claims and Extension of Time (EOT) requests.",
             "Automated cross-referencing of project baseline schedules, daily site logs, and FIDIC contractual clauses for rapid dispute resolution.",
             "Provides automated claim vulnerability scoring and delay attribution matrices to assist claims consultants and project managers."
         ]),
        ("Government Primary School KPK — Complete 3D/4D/5D BIM Project",
         "BIM PAK / BRAINS Institute Peshawar",
         "BIM & VDC Infrastructure Project",
         [
             "Led complete BIM lifecycle modeling in Revit (Architectural & Structural coordination) for an educational facility in KPK.",
             "Executed 4D construction sequencing in Navisworks Manage linked with Primavera P6 5D cost and 12-month schedule analysis.",
             "Produced high-fidelity Twinmotion photorealistic walkthroughs and BOQs, reducing on-site rework by 15%."
         ])
    ]

    for p_title, p_link, p_type, bullets in ai_projects:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(4)
        p_head.paragraph_format.space_after = Pt(2)
        p_head.paragraph_format.keep_with_next = True

        r_ht = p_head.add_run(p_title)
        r_ht.font.bold = True
        r_ht.font.size = Pt(10)
        r_ht.font.color.rgb = RGBColor(15, 23, 42)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(3)
        p_sub.paragraph_format.keep_with_next = True
        
        r_sub1 = p_sub.add_run(f"{p_type}   |   {p_link}")
        r_sub1.font.size = Pt(8.5)
        r_sub1.font.italic = True
        r_sub1.font.color.rgb = RGBColor(100, 116, 139)

        for b in bullets:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_after = Pt(2)
            pb.paragraph_format.left_indent = Inches(0.2)
            rb = pb.add_run("•  " + b)
            rb.font.size = Pt(9)

    # ==========================================
    # PROFESSIONAL EXPERIENCE
    # ==========================================
    add_heading_styled(doc, "PROFESSIONAL EXPERIENCE (15+ YEARS)")

    experiences = [
        ("MANAGER ESTATE (TECHNICAL)", "REDAMCO — MINISTRY OF RAILWAYS", "Rawalpindi / Islamabad", "Dec 2025 – Present",
         "Public Sector Corporate & Asset Management",
         [
             "Analyze technical feasibilities, execution strategies, and structural/infrastructural appraisals for major Ministry of Railways estate projects.",
             "Assist the Senior Manager Estate and Business Development in strategic land monetization, commercial proposals, joint ventures, and revenue optimization.",
             "Perform technical risk assessments, lease compliance audits, and engineering evaluations on railway land assets across Pakistan.",
             "Implement digital project monitoring, GIS asset mapping, and BIM workflows to streamline real estate asset performance and commercial delivery."
         ]),
        ("INSPECTOR OF WORKS / INFRASTRUCTURE ENGINEER", "PAKISTAN RAILWAYS", "Rawalpindi, Pakistan", "May 2018 – 09th Dec 2025 (91 Months)",
         "Public Sector Client / Infrastructure",
         [
             "Planned, executed, and monitored major railway infrastructure projects including track renewals, station rehabilitation, and bridge maintenance.",
             "Conducted detailed land and topographic surveys using GIS systems, removed ROW encroachments, and managed railway asset safety.",
             "Prepared cost estimates, engineering budgets, technical drawings, structural plans, and liaised with external departments for statutory approvals.",
             "Coordinated multi-disciplinary traffic and civil engineering operations to ensure zero-disruption train transit during site execution."
         ]),
        ("SITE ENGINEER (CIVIL)", "SHAH ZAMAN CONSTRUCTIONS PVT LTD", "Rawalpindi, Pakistan", "Mar 2017 – Jan 2018 (10 Months)",
         "DHA Emaar Housing Society Infrastructure Project",
         [
             "Prepared and submitted Interim Payment Certificates (IPCs), variation orders, and extension of time (EOT) claims.",
             "Managed material procurement, Bar Bending Schedules (BBS), material demand lists, and BOQ compliance.",
             "Maintained detailed project execution schedules using Primavera P6, ensuring progress monitoring within budget constraints.",
             "Oversaw on-site quality control testing, verified subcontractor billings against completed works, and maintained contract documents."
         ]),
        ("SITE ENGINEER (CIVIL)", "CONSTRUCTION INVADAR PVT LTD", "Rawalpindi, Pakistan", "Oct 2013 – Jul 2016 (33 Months)",
         "Housing Directorate GHQ (Askari 5, 11 & 14 Projects)",
         [
             "Supervised construction of 2xG+4 residential flats, luxury villas, internal road networks, and infrastructure works.",
             "Managed comprehensive site documentation including IPCs, Running Account Bills (RARs), deviation orders, and BBS.",
             "Enforced quality compliance through rigorous concrete testing, structural inspection, and batching plant coordination.",
             "Evaluated subcontractor performance and maintained progress schedules to achieve targeted milestones."
         ]),
        ("SITE ENGINEER (CIVIL)", "AL RAAK INTERNATIONAL LLC", "Muscat & Sohar, OMAN", "May 2011 – Sep 2013 (28 Months)",
         "Sohar Steel Mill Structure & Residential Villas (GCC Experience)",
         [
             "Executed heavy frame structure building works for Sohar Steel Mill and multiple luxury residential villas in Oman.",
             "Oversee RCC works including slope protection, access ramps, concrete pavements, and precast element erection.",
             "Managed material procurement, layout marking, concrete batching plant delivery, and earthwork grading precision."
         ]),
        ("SUB ENGINEER (CIVIL)", "AL SUBHAN CORPORATION", "Rawalpindi, Pakistan", "May 2010 – Jan 2011 (8 Months)",
         "G+3 Residential Flats near CMH Rawalpindi",
         [
             "Supervised structural concrete & steel works, site layout surveys, labor allocation, and material inventory.",
             "Prepared detailed work schedules, conducted joint inspections with consultants, and processed subcontractor invoices."
         ]),
        ("SITE SUPERVISOR (CIVIL)", "IKSA CONSTRUCTIONS", "Rawalpindi, Pakistan", "Jan 2010 – May 2010 (4 Months)",
         "Multistorey Ghakkar Plaza Commercial Project",
         [
             "Supervised site layout marking, concrete pouring, steel reinforcement verification, and daily workforce output."
         ])
    ]

    for role, company, loc, dates, proj_info, bullets in experiences:
        p_ehead = doc.add_paragraph()
        p_ehead.paragraph_format.space_before = Pt(6)
        p_ehead.paragraph_format.space_after = Pt(1)
        p_ehead.paragraph_format.keep_with_next = True

        r_r = p_ehead.add_run(role)
        r_r.font.bold = True
        r_r.font.size = Pt(10)
        r_r.font.color.rgb = RGBColor(15, 23, 42)

        r_c = p_ehead.add_run(f"  |  {company}")
        r_c.font.bold = True
        r_c.font.size = Pt(10)
        r_c.font.color.rgb = RGBColor(13, 148, 136)

        p_esub = doc.add_paragraph()
        p_esub.paragraph_format.space_after = Pt(3)
        p_esub.paragraph_format.keep_with_next = True
        
        r_es = p_esub.add_run(f"{loc}   |   {dates}   |   Scope: {proj_info}")
        r_es.font.size = Pt(8.5)
        r_es.font.italic = True
        r_es.font.color.rgb = RGBColor(100, 116, 139)

        for b in bullets:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_after = Pt(2)
            pb.paragraph_format.left_indent = Inches(0.2)
            rb = pb.add_run("•  " + b)
            rb.font.size = Pt(9)

    # ==========================================
    # EDUCATION & PROFESSIONAL QUALIFICATIONS
    # ==========================================
    add_heading_styled(doc, "EDUCATION & ACADEMIC QUALIFICATIONS")

    edu_list = [
        ("Applied AI & LLM Engineering", "Edersity Islamabad", "2025 – 2026", "Specialization in AI Agentic Workflows, LLM RAG Architecture, and AI-driven Automation."),
        ("Post Graduate Diploma (PGD) in BIM & Construction Management", "BRAINS Institute Peshawar", "2025", "Specialization in Revit, Navisworks Manage (4D/5D), ISO 19650, Primavera P6, Civil 3D & VDC."),
        ("BS (Honors) in Civil Technology", "Preston University Kohat, Islamabad Campus", "2019", "3.6 GPA | RCC Design, Steel Structures, Highway & Tunnel Engineering, Theory of Structures."),
        ("Bachelor of Arts (Journalism & Education)", "University of Sargodha", "2015", "Graduated in Mass Communication, Media & Educational Methodologies."),
        ("Diploma of Associate Engineer (DAE Civil)", "Government College of Technology Rasul", "2009", "70% Marks | Quantity Surveying, Soil Mechanics, Construction Execution & RCC Design."),
        ("Secondary School Certificate (SSC Science)", "APS & College Dhamial Camp Rawalpindi", "2006", "72% Marks | Physics, Chemistry, Mathematics & Computer Science.")
    ]

    for deg, inst, yr, desc in edu_list:
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(3)
        p_edu.paragraph_format.space_after = Pt(1)
        p_edu.paragraph_format.keep_with_next = True

        r_d = p_edu.add_run(deg)
        r_d.font.bold = True
        r_d.font.size = Pt(9.5)
        r_d.font.color.rgb = RGBColor(15, 23, 42)

        r_i = p_edu.add_run(f" — {inst} ({yr})")
        r_i.font.size = Pt(9)
        r_i.font.color.rgb = RGBColor(13, 148, 136)

        p_edesc = doc.add_paragraph()
        p_edesc.paragraph_format.space_after = Pt(3)
        p_edesc.paragraph_format.left_indent = Inches(0.2)
        r_dd = p_edesc.add_run(desc)
        r_dd.font.size = Pt(8.5)
        r_dd.font.color.rgb = RGBColor(71, 85, 105)

    # ==========================================
    # CERTIFICATIONS & ACCREDITATIONS
    # ==========================================
    add_heading_styled(doc, "CERTIFICATIONS & PROFESSIONAL CREDENTIALS")

    certs = [
        "Applied AI & LLM Engineer — Edersity Islamabad (2025/2026)",
        "Plannerly Certified Information Manager (ISO 19650 Level 1, 2, 3 Expert & Certification Ready)",
        "Mastering BIM with ISO 19650 — BIMPACC (2025)",
        "CanBIM Foundations Software Skills - Proficient (S4) — CanBIM / Skill Development Council Canada (2025)",
        "Autodesk Civil 3D Certification — Skill Development Council Canada (2025)",
        "Graduate Engineering Technologist — National Technology Council (NTC) Islamabad (2019)",
        "Diploma in Construction Management with Primavera P6 — Primavera Institute (2022)",
        "Revit Architecture Certification — PNY Trainings (2024)",
        "AutoCAD 2D & 3D Certification — E-Hunar (2024)",
        "Full-Stack Web Development Certification — Talha Tariq LMS (2025)",
        "AI Agentic Workflows & Automation — N8N Skool Community (2025)"
    ]

    for c in certs:
        pc = doc.add_paragraph()
        pc.paragraph_format.space_after = Pt(2)
        pc.paragraph_format.left_indent = Inches(0.2)
        rc = pc.add_run("✓  " + c)
        rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    print(f"Updated CV DOCX successfully generated at: {output_path}")

if __name__ == "__main__":
    out_dir = r"C:\Users\DELL\.gemini\antigravity\scratch\waqas-cv-portfolio"
    docx_path = os.path.join(out_dir, "Engr_Muhammad_Waqas_CV.docx")
    photo_path = os.path.join(out_dir, "profile_photo.png")
    create_cv_docx(docx_path, photo_path)
