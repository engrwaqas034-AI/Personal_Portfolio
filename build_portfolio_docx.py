import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="14" w:space="4" w:color="0D9488"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def create_portfolio_docx(output_path, img_path):
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Base Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(51, 65, 85)

    # ==========================================
    # HEADER SECTION WITH PHOTO
    # ==========================================
    tbl_header = doc.add_table(rows=1, cols=2)
    tbl_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_header.autofit = False

    cell_text = tbl_header.rows[0].cells[0]
    cell_img = tbl_header.rows[0].cells[1]

    cell_text.width = Inches(5.8)
    cell_img.width = Inches(1.3)

    set_cell_margins(cell_text, top=0, bottom=0, left=0, right=50)
    set_cell_margins(cell_img, top=0, bottom=0, left=50, right=0)
    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_name = cell_text.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run("MUHAMMAD WAQAS")
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(15, 23, 42)

    p_title = cell_text.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("CIVIL ENGINEER  |  BIM & VDC SPECIALIST  |  APPLIED AI & LLM ENGINEER")
    r_title.font.size = Pt(9.5)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(13, 148, 136)

    p_contact = cell_text.add_paragraph()
    p_contact.paragraph_format.space_after = Pt(6)
    r_contact = p_contact.add_run(
        "Rawalpindi, Pakistan  |  +92 345 5862998  |  engrwaqas034@gmail.com\n"
        "Live Portfolio: engrwaqas034-ai.github.io/Personal_Portfolio  |  LinkedIn: linkedin.com/in/muhammad-waqas-27b232214  |  GitHub: github.com/engrwaqas034-AI"
    )
    r_contact.font.size = Pt(8.5)
    r_contact.font.color.rgb = RGBColor(71, 85, 105)

    if os.path.exists(img_path):
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(img_path, width=Inches(1.25))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ==========================================
    # SECTION 1: EXECUTIVE VISION
    # ==========================================
    add_heading_styled(doc, "1. EXECUTIVE PROFILE & HYBRID ENGINEERING VISION")
    
    p_vis = doc.add_paragraph()
    p_vis.paragraph_format.space_after = Pt(8)
    p_vis.paragraph_format.line_spacing = 1.15
    p_vis.add_run(
        "Muhammad Waqas represents a new paradigm in the Built Environment: a Civil Engineer, Manager Estate (Technical) at REDAMCO (Ministry of Railways) "
        "and Senior Infrastructure Specialist with 15+ years of hands-on field, infrastructure, and project management leadership who has systematically integrated "
        "Building Information Modeling (BIM / VDC), ISO 19650 Information Management, and Applied Artificial Intelligence (LLM & Agentic Workflows). "
        "With a proven background bridging corporate railway asset monetization, client-side megaprojects (Pakistan Railways), GCC industrial mega-structures "
        "(Sohar Steel Mill Oman), and multi-story residential/commercial developments, Waqas designs, executes, and automates modern construction workflows."
    )

    # ==========================================
    # SECTION 2: AI & DIGITAL AEC INNOVATION PROJECTS
    # ==========================================
    add_heading_styled(doc, "2. FEATURED AI & DIGITAL AEC INNOVATION PROJECTS")

    projects = [
        ("PROJECT 01: AI-Architect — Generative AI Layout & BIM Engine",
         "Domain: AI Engineering, Generative Architecture & CAD Automation",
         "GitHub Repository: https://github.com/engrwaqas034-AI/AI-Architect",
         [
             "Concept & Innovation: Engineered a cutting-edge generative AI platform that automates 2D/3D architectural floorplan generation based on plot dimensions, setback rules, and functional zoning requirements.",
             "Technical Stack: Python, Custom CSP (Constraint Satisfaction Problem) Solver, SVG Parametric Engine, RAG Knowledge Base, HTML5/CSS3 Interactive UI.",
             "Key Functionalities: Calculates zoning envelopes, auto-places living/dining/bedroom spaces with zero overlapping, verifies setback compliance, and exports standard DXF/BIM coordinate structures.",
             "Impact: Reduces preliminary architectural floorplan drafting time from days to seconds while eliminating spatial design errors."
         ]),
        ("PROJECT 02: Construction_Claim_Analyzer — LLM & RAG Delay Claim Intelligence",
         "Domain: AI Engineering, Claims & Dispute Resolution, Construction Law",
         "GitHub Repository: https://github.com/engrwaqas034-AI/Construction_Claim_Analyzer",
         [
             "Concept & Innovation: Developed an intelligent LLM & RAG (Retrieval-Augmented Generation) analytical system specialized in evaluating construction delay claims, Extension of Time (EOT) applications, and contractual disputes.",
             "Contractual Knowledge Base: Embedded FIDIC standard contracts, project baseline schedules (Primavera P6), daily site reports, and variation orders into a searchable vector database.",
             "Automated Analysis: Performs instant delay attribution, critical path disruption checks, and contract clause compliance verification, outputting structured legal/engineering risk scores.",
             "Impact: Empowers claims consultants and project managers to audit complex contractual claims in minutes rather than weeks."
         ]),
        ("PROJECT 03: AI BOQ & Quantity Surveying Generator",
         "Domain: AI Engineering, Quantity Surveying & Cost Estimation",
         "Repository: Integrated AEC AI Tool Pipeline",
         [
             "Concept & Innovation: Built an automated AI-driven pipeline for extracting Bill of Quantities (BOQ), material demands, and structural steel/concrete quantities directly from architectural inputs.",
             "Technical Features: Parses structural parameters, calculates concrete volumes, bar bending schedules (BBS), and formwork areas, linking directly with cost databases.",
             "Impact: Eliminates manual take-off errors, drastically accelerating pre-construction estimation and tender preparation."
         ])
    ]

    for title, domain, repo, bullets in projects:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.keep_with_next = True
        r_t = p_t.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(15, 23, 42)

        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(4)
        p_meta.paragraph_format.keep_with_next = True
        r_m = p_meta.add_run(f"{domain}  |  {repo}")
        r_m.font.size = Pt(8.5)
        r_m.font.italic = True
        r_m.font.color.rgb = RGBColor(13, 148, 136)

        for b in bullets:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_after = Pt(2)
            pb.paragraph_format.left_indent = Inches(0.2)
            rb = pb.add_run("▪  " + b)
            rb.font.size = Pt(9)

    # ==========================================
    # SECTION 3: FEATURED BIM & VDC PROJECTS
    # ==========================================
    add_heading_styled(doc, "3. FEATURED BIM & VDC (BUILDING INFORMATION MODELING) PROJECTS")

    bim_projects = [
        ("PROJECT 04: Construction of Government Primary School KPK (Complete BIM Delivery)",
         "Location: Khyber Pakhtunkhwa, Pakistan  |  Role: Lead BIM Engineer  |  Tools: Revit, Navisworks Manage, Primavera P6, Twinmotion",
         [
             "Integrated 3D BIM Modeling: Created high-detail architectural and structural models in Autodesk Revit, ensuring full multi-disciplinary clash detection.",
             "4D Construction Simulation: Linked 3D Revit models with Primavera P6 schedules in Navisworks Manage for step-by-step visual construction sequencing.",
             "5D Cost & Quantity Tracking: Generated automated BOQs, material schedules, and progress cost tracking directly from the BIM database.",
             "Visualization & Impact: Rendered photorealistic 3D walkthroughs in Twinmotion, achieving a +15% increase in site construction efficiency."
         ]),
        ("PROJECT 05: Modern Residential Villa (3D BIM & Real-Time Visualization)",
         "Location: Rawalpindi / Islamabad  |  Role: BIM Designer & Coordinator  |  Tools: Revit, Twinmotion, AutoCAD",
         [
             "3D Spatial Design: Developed complete 3D BIM models emphasizing space optimization, natural illumination, and modern architectural aesthetics.",
             "Visualization & VR Walkthroughs: Rendered real-time client presentations in Twinmotion featuring realistic lighting, material textures, and spatial flows.",
             "Construction Documentation: Extracted precise 2D architectural plans, elevations, sections, and structural reinforcement details."
         ]),
        ("PROJECT 06: Civil 3D Road Infrastructure & Earthwork Modeling",
         "Location: Infrastructure Sector  |  Role: Civil Infrastructure Specialist  |  Tools: Autodesk Civil 3D",
         [
             "Surface & Topography: Generated accurate 3D terrain surfaces from existing survey point clouds and GIS datasets.",
             "Corridor Design: Designed road alignments, vertical profiles, assemblies (lanes, shoulders, ditches), and 3D corridors.",
             "Earthwork Optimization: Calculated precise cut-and-fill volumes, optimizing grading design to minimize earthwork costs."
         ])
    ]

    for title, meta, bullets in bim_projects:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.keep_with_next = True
        r_t = p_t.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(15, 23, 42)

        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_after = Pt(4)
        p_m.paragraph_format.keep_with_next = True
        r_m = p_m.add_run(meta)
        r_m.font.size = Pt(8.5)
        r_m.font.italic = True
        r_m.font.color.rgb = RGBColor(13, 148, 136)

        for b in bullets:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_after = Pt(2)
            pb.paragraph_format.left_indent = Inches(0.2)
            rb = pb.add_run("▪  " + b)
            rb.font.size = Pt(9)

    # ==========================================
    # SECTION 4: MAJOR CIVIL INFRASTRUCTURE EXPERIENCE
    # ==========================================
    add_heading_styled(doc, "4. MAJOR CIVIL & INFRASTRUCTURE PROJECTS (15+ YEARS FIELD EXPERIENCE)")

    field_projects = [
        ("REDAMCO (Railway Estate Development & Marketing Company) — Ministry of Railways (Dec 2025 – Present)",
         "Role: Manager Estate (Technical)  |  Public Sector Corporate & Asset Management",
         [
             "Performs technical appraisals, execution feasibility reviews, and engineering evaluations for Ministry of Railways real estate developments.",
             "Assists Senior Manager Estate and Business Development in land monetization strategies, commercial joint ventures, and revenue optimization.",
             "Conducts structural/infrastructural risk assessments and lease compliance audits on commercial railway assets across Pakistan."
         ]),
        ("Pakistan Railways Infrastructure & Railway Network Projects (May 2018 – 09th Dec 2025)",
         "Role: Inspector of Works / Infrastructure Engineer  |  Client: Pakistan Railways (Public Sector)",
         [
             "Oversees maintenance, track renewal, station rehabilitation, and bridge safety across major railway corridors.",
             "Utilizes GIS land surveying to clear right-of-way (ROW) encroachments and map railway land boundaries.",
             "Prepares structural designs, BOQs, technical specifications, and liaises with government agencies for project sanctions."
         ]),
        ("Housing Directorate GHQ — Askari 5, 11 & 14 Projects (Oct 2013 – Jul 2016)",
         "Role: Site Engineer (Civil)  |  Contractor: Construction Invadar Pvt Ltd",
         [
             "Supervised construction of 2xG+4 residential flats, luxury villas, internal road networks, drainage, and utilities.",
             "Managed IPCs, Running Account Bills (RARs), deviation orders, Bar Bending Schedules (BBS), and subcontractor payments."
         ]),
        ("Sohar Steel Mill Structure & Residential Villas, OMAN (May 2011 – Sep 2013)",
         "Role: Site Engineer (Civil)  |  Company: Al Raak International LLC (Muscat & Sohar, Oman)",
         [
             "Executed heavy structural steel mill framing, precast concrete erection, slope protection, and access ramps in Sohar, Oman.",
             "Supervised batching plant concrete deliveries, site grading, paving, and measurement verification."
         ]),
        ("DHA Emaar Housing Society Infrastructure Project (Mar 2017 – Jan 2018)",
         "Role: Site Engineer (Civil)  |  Contractor: Shah Zaman Constructions Pvt Ltd",
         [
             "Managed site execution, baseline scheduling in Primavera P6, material procurement, BBS, and interim payment claims."
         ])
    ]

    for title, meta, bullets in field_projects:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.keep_with_next = True
        r_t = p_t.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)

        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_after = Pt(4)
        p_m.paragraph_format.keep_with_next = True
        r_m = p_m.add_run(meta)
        r_m.font.size = Pt(8.5)
        r_m.font.italic = True
        r_m.font.color.rgb = RGBColor(13, 148, 136)

        for b in bullets:
            pb = doc.add_paragraph()
            pb.paragraph_format.space_after = Pt(2)
            pb.paragraph_format.left_indent = Inches(0.2)
            rb = pb.add_run("▪  " + b)
            rb.font.size = Pt(9)

    doc.save(output_path)
    print(f"Updated Portfolio DOCX successfully generated at: {output_path}")

if __name__ == "__main__":
    out_dir = r"C:\Users\DELL\.gemini\antigravity\scratch\waqas-cv-portfolio"
    doc_path = os.path.join(out_dir, "Engr_Muhammad_Waqas_Portfolio.docx")
    photo_path = os.path.join(out_dir, "profile_photo.png")
    create_portfolio_docx(doc_path, photo_path)
